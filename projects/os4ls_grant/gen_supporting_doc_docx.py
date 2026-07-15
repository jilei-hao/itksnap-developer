#!/usr/bin/env python3
"""Regenerate OS4LS supporting-document .docx from the revised markdown,
re-embedding the figures. Produces a reviewable, page-budget-conscious doc."""
import re, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import os
ROOT=os.path.dirname(os.path.abspath(__file__))

# figure files (full-width for legibility; sizes tuned to keep the doc <=4 pages)
FIG = {
 "Fig. 1": (f"{ROOT}/supporting_docs/fig_ui-screenshots.png", 6.6, ""),          # UI capabilities (R01 fig1)
 "Fig. 2": (f"{ROOT}/supporting_docs/fig_ai-placenta.png", 6.2, ""),             # interactive AI (R01 fig3)
 "Fig. 3": (f"{ROOT}/supporting_docs/figures/fig_workflow.png", 6.6, ""),        # workflow (wide/short)
 "Fig. 4": (f"{ROOT}/supporting_docs/figures/fig_architecture.png", 6.6, ""),    # architecture (tall)
}

md=open(f"{ROOT}/supporting_document_draft.md").read()
# Drop the trailing production-notes block. Cut from the FIRST comment-open to end of file
# (notes are always the trailing block) — robust even if the note text mentions the literal
# comment markers, which would defeat a non-greedy <!--...--> match.
md=re.sub(r'<!--.*\Z','',md,flags=re.DOTALL).strip()

doc=Document()
sec=doc.sections[0]
sec.page_height=Inches(11); sec.page_width=Inches(8.5)
for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec,m,Inches(0.9))
normal=doc.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(10.5)
normal.paragraph_format.space_after=Pt(4); normal.paragraph_format.line_spacing=1.03

HEAD=RGBColor(0xC0,0x50,0x2B)  # muted orange like the fund template

def add_runs(p, text):
    """Parse **bold**, *italic*, `code` into runs."""
    tokens=re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for tok in tokens:
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'):
            r=p.add_run(tok[2:-2]); r.bold=True
        elif tok.startswith('*') and tok.endswith('*'):
            r=p.add_run(tok[1:-1]); r.italic=True
        elif tok.startswith('`') and tok.endswith('`'):
            r=p.add_run(tok[1:-1]); r.font.name='Consolas'; r.font.size=Pt(9.5)
        else:
            p.add_run(tok)

def add_image(key):
    path,width,note=FIG[key]
    if os.path.exists(path):
        pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before=Pt(2); pp.paragraph_format.space_after=Pt(2)
        pp.add_run().add_picture(path, width=Inches(width))
    if note:
        n=doc.add_paragraph(); n.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=n.add_run(note.strip()); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x88,0x88,0x88)

lines=md.split('\n')
i=0; buf=[]; pending_fig=[None]
def flush():
    global buf
    t=' '.join(buf).strip()
    if t:
        p=doc.add_paragraph(); add_runs(p,t)
        p.paragraph_format.space_after=Pt(3)
        if pending_fig[0]:
            p.paragraph_format.keep_with_next=True   # caption stays with its figure
            add_image(pending_fig[0]); pending_fig[0]=None
    buf=[]

def fig_key(s):
    m=re.match(r'\*{0,2}(Fig\. \d)', s)
    return m.group(1) if (m and m.group(1) in FIG) else None

while i < len(lines):
    ln=lines[i]; s=ln.strip()
    # table block
    if s.startswith('|'):
        flush()
        rows=[]
        while i < len(lines) and lines[i].strip().startswith('|'):
            cells=[c.strip() for c in lines[i].strip().strip('|').split('|')]
            if not set(''.join(cells))<=set('-: '):
                rows.append(cells)
            i+=1
        if rows:
            ncol=len(rows[0]); tbl=doc.add_table(rows=0,cols=ncol); tbl.style='Table Grid'
            for ri,row in enumerate(rows):
                cellrow=tbl.add_row().cells
                for ci,c in enumerate(row[:ncol]):
                    cellrow[ci].text=''
                    para=cellrow[ci].paragraphs[0]; para.paragraph_format.space_after=Pt(1)
                    add_runs(para, c)
                    for rr in para.runs:
                        rr.font.size=Pt(9)
                        if ri==0: rr.bold=True
        continue
    if s.startswith('# '):
        flush()
        h=doc.add_paragraph(); h.paragraph_format.space_before=Pt(2); h.paragraph_format.space_after=Pt(6)
        r=h.add_run(s[2:]); r.bold=True; r.font.size=Pt(14)
        i+=1; continue
    if s.startswith('## '):
        flush()
        h=doc.add_paragraph(); h.paragraph_format.space_before=Pt(8); h.paragraph_format.space_after=Pt(3)
        r=h.add_run(s[3:]); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=HEAD
        i+=1; continue
    if s.startswith('- '):
        flush()
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
        add_runs(p, s[2:])
        i+=1; continue
    # numbered reference item: flush previous, start a new hanging-indent paragraph
    if re.match(r'^\d+\.\s', s) and not s.startswith('## '):
        flush()
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
        p.paragraph_format.left_indent=Inches(0.28); p.paragraph_format.first_line_indent=Inches(-0.28)
        # accumulate the FULL reference (incl. wrapped continuation lines) before parsing,
        # so *italic* spans that cross a line wrap are matched as a pair
        ref_text=s; i+=1
        while i < len(lines):
            cs=lines[i].strip()
            if cs=='' or cs.startswith(('#','|','- ')) or re.match(r'^\d+\.\s',cs): break
            ref_text += ' '+cs; i+=1
        add_runs(p, ref_text)
        continue
    if s=='' or s=='---':
        flush(); i+=1; continue
    buf.append(s)
    # figure caption -> mark pending so the image is inserted AFTER the full caption paragraph
    if fig_key(s):
        pending_fig[0]=fig_key(s)
    i+=1
flush()

out=f"{ROOT}/OS4LS_Supporting_Document_ITK-SNAP.docx"
doc.save(out)
print("saved", out)
